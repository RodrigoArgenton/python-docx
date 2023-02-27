from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def tableStyle(s):
    '''
    Editar a o estilo da tabela.

    s: indice do estilo que se encontra na lista 'style'.
    return: retorna template escolhido ou o padrão, caso o valor inserido esteja fora da lista.
    '''
    style = (
        'Table Grid',        # [0]
        'Light Shading',     # [1]
        'Light List',        # [2]
        'Medium Shading 1',  # [3]
        'Medium Shading 2',  # [4]
        'Medium List 1',     # [5]
        'Medium List 2',     # [6]
        'Medium Grid 1',     # [7]
        'Medium Grid 2',     # [8]
        'Medium Grid 3',     # [9]
    )
    # laço de verificação e retorno do estilo da tabela
    for i, v in enumerate(style):
        if s == i:
            return v
    return 'Table Grid'


def newTable(styleTable=0, condHeaderCenter=False, condTableCenter=False,):
    '''
    Criar e editar estilo de uma tabela fixa, essa função tem como objetivo receber 3 parametros, sendo eles:

    styleTable: recebe um valor de 0 até 10. Obs: caso o valor inserido não esteja dentro do informado, será utililizado o valor 0, ou seja, o estilo Table Grid.
    condHeaderCenter: Centraliza a primeira linha da tabela.
    CondTableCenter: Centraliza toda a tabela.
    return: retorna um documento ja salvo com a lista e informações criadas.
    '''

    # Lista fixa
    items = (
        (1, 'rua barão, 987', 'Marcio da silva'),
        (2, 'rua trad, 67', 'José silva'),
        (3, 'rua marcinho, 09', 'Maria silva'),
    )

    # função para criar ou abrir um documento caso seja passado um parametro dentro de Document()
    document = Document()

    # Criar tabela com 1 linha e 3 colunas
    table = document.add_table(rows=1, cols=3)

    # Editar estilo da tabela com a função tableStyle
    table.style = tableStyle(styleTable)

    # Preencher informações no cabeçalho da tabela
    head_cell = table.rows[0].cells
    head_cell[0].text = 'Number'
    head_cell[1].text = 'Endereço'
    head_cell[2].text = 'Nome completo'

    # laço com o objetivo de inserir valores da lista 'item' na tabela. o mesmo tem a fnção de inserir uma nova linha
    for item in items:
        cells = table.add_row().cells
        cells[0].text = str(item[0])
        cells[1].text = item[1]
        cells[2].text = item[2]

    # Condicional com o objetivo de colocar ou não, os valor centralizado no cabeçalho
    if condHeaderCenter is True:
        for cell in table.rows[0].cells:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Condicional com o objetivo de colocar ou não, os valor centralizado na tabela
    if condTableCenter is True:
        # Percorrer todas as células da tabela
        for row in table.rows:
            for cell in row.cells:
                # Define o alinhamento horizontal da célula como centralizado
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # salvar documento
    document.save('teste.docx')
